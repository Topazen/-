from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Створити новий документ
doc = Document()

# Налаштувати поля документа
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# Заголовок
title = doc.add_paragraph("ПРОТОКОЛ РОБОТИ ПРОГРАМИ")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].bold = True
title.runs[0].font.size = Pt(16)
title.runs[0].font.name = 'Times New Roman'

subtitle = doc.add_paragraph("Система керування торговим автоматом")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.name = 'Times New Roman'

doc.add_paragraph()

# Розділ 1: Запуск системи
heading1 = doc.add_paragraph("1. Запуск системи")
heading1.runs[0].bold = True
heading1.runs[0].font.size = Pt(14)
heading1.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph("При запуску програми користувачу відображається головне меню торгового автомата з переліком доступних товарів та опцій.")
p.runs[0].font.size = Pt(14)
p.runs[0].font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Додати перший скріншот
img_path1 = r"C:\Users\nelox\.gemini\antigravity\brain\c686f7b8-872e-4a7f-ad27-6f83ba57b924\vending_start_screen_1764038866097.png"
if os.path.exists(img_path1):
    doc.add_picture(img_path1, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

caption = doc.add_paragraph("Рисунок 1. Головний екран програми")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.runs[0].font.size = Pt(12)
caption.runs[0].font.name = 'Times New Roman'
caption.runs[0].italic = True

doc.add_paragraph()

# Розділ 2: Вибір товару
heading2 = doc.add_paragraph("2. Сценарій 1: Успішна купівля товару")
heading2.runs[0].bold = True
heading2.runs[0].font.size = Pt(14)
heading2.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph("Кроки виконання:")
p.runs[0].font.size = Pt(14)
p.runs[0].font.name = 'Times New Roman'

steps = [
    "Користувач обирає опцію '1. Обрати товар'",
    "Вводить ID товару (наприклад, 1 для Coca-Cola)",
    "Обирає опцію '2. Внести кошти'",
    "Вводить необхідну суму (₴15.00)",
    "Система видає товар і повертає решту (якщо є)"
]

for step in steps:
    p = doc.add_paragraph(step, style='List Bullet')
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph("Очікуваний результат: Товар успішно видано, решта повернена користувачу.")
p.runs[0].font.size = Pt(14)
p.runs[0].font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Додати другий скріншот
img_path2 = r"C:\Users\nelox\.gemini\antigravity\brain\c686f7b8-872e-4a7f-ad27-6f83ba57b924\vending_purchase_success_1764038879565.png"
if os.path.exists(img_path2):
    doc.add_picture(img_path2, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

caption = doc.add_paragraph("Рисунок 2. Процес купівлі товару")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.runs[0].font.size = Pt(12)
caption.runs[0].font.name = 'Times New Roman'
caption.runs[0].italic = True

doc.add_paragraph()

# Розділ 3: Скасування транзакції
heading3 = doc.add_paragraph("3. Сценарій 2: Скасування транзакції")
heading3.runs[0].bold = True
heading3.runs[0].font.size = Pt(14)
heading3.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph("Кроки виконання:")
p.runs[0].font.size = Pt(14)
p.runs[0].font.name = 'Times New Roman'

steps2 = [
    "Користувач вносить гроші (опція 2)",
    "Передумує купувати",
    "Обирає опцію '3. Скасувати транзакцію'",
    "Система повертає внесені кошти"
]

for step in steps2:
    p = doc.add_paragraph(step, style='List Bullet')
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph("Очікуваний результат: Транзакція скасована, гроші повернені, стан автомата скинуто до початкового.")
p.runs[0].font.size = Pt(14)
p.runs[0].font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# Розділ 4: Недостатньо коштів
heading4 = doc.add_paragraph("4. Сценарій 3: Недостатньо коштів")
heading4.runs[0].bold = True
heading4.runs[0].font.size = Pt(14)
heading4.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph("Кроки виконання:")
p.runs[0].font.size = Pt(14)
p.runs[0].font.name = 'Times New Roman'

steps3 = [
    "Користувач обирає товар Coca-Cola (₴15.00)",
    "Вносить лише ₴10.00",
    "Намагається завершити покупку"
]

for step in steps3:
    p = doc.add_paragraph(step, style='List Bullet')
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph("Очікуваний результат: Система виводить повідомлення 'Недостатньо коштів. Внесено: ₴10.00, Потрібно: ₴15.00'")
p.runs[0].font.size = Pt(14)
p.runs[0].font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# Розділ 5: Режим адміністратора
heading5 = doc.add_paragraph("5. Сценарій 4: Адміністрування")
heading5.runs[0].bold = True
heading5.runs[0].font.size = Pt(14)
heading5.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph("Кроки виконання:")
p.runs[0].font.size = Pt(14)
p.runs[0].font.name = 'Times New Roman'

steps4 = [
    "Користувач обирає опцію '4. Адмін-меню'",
    "Вводить пароль 'admin'",
    "Обирає '1. Поповнити товар'",
    "Вводить ID товару (наприклад, 1)",
    "Вводить кількість (наприклад, 5)",
    "Товар додається до інвентарю"
]

for step in steps4:
    p = doc.add_paragraph(step, style='List Bullet')
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph("Очікуваний результат: Інвентар успішно поповнено. Логується повідомлення 'Адмін поповнив товар 1 на 5 шт.'")
p.runs[0].font.size = Pt(14)
p.runs[0].font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# Розділ 6: Результати тестування
heading6 = doc.add_paragraph("6. Результати тестування")
heading6.runs[0].bold = True
heading6.runs[0].font.size = Pt(14)
heading6.runs[0].font.name = 'Times New Roman'

# Створити таблицю
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'

# Заголовок таблиці
header_cells = table.rows[0].cells
header_cells[0].text = '№'
header_cells[1].text = 'Сценарій'
header_cells[2].text = 'Результат'

for cell in header_cells:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(14)
    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'

# Рядки даних
scenarios = [
    ('1', 'Успішна купівля товару', 'Пройдено ✓'),
    ('2', 'Скасування транзакції', 'Пройдено ✓'),
    ('3', 'Недостатньо коштів', 'Пройдено ✓'),
    ('4', 'Поповнення інвентарю', 'Пройдено ✓')
]

for i, (num, scenario, result) in enumerate(scenarios, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = num
    row_cells[1].text = scenario
    row_cells[2].text = result
    
    for cell in row_cells:
        cell.paragraphs[0].runs[0].font.size = Pt(14)
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'

doc.add_paragraph()

# Висновки
heading7 = doc.add_paragraph("7. Висновки")
heading7.runs[0].bold = True
heading7.runs[0].font.size = Pt(14)
heading7.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph("Програмна система торгового автомата успішно виконує всі заявлені функції:")
p.runs[0].font.size = Pt(14)
p.runs[0].font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

conclusions = [
    "Відображення доступних товарів та їх кількості",
    "Прийом коштів та обчислення решти",
    "Видача товарів при достатній кількості коштів",
    "Скасування транзакцій з поверненням грошей",
    "Адміністративні функції (поповнення інвентарю)",
    "Коректна обробка помилок (недостатньо коштів, відсутність товару)"
]

for conclusion in conclusions:
    p = doc.add_paragraph(conclusion, style='List Bullet')
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph("Система відповідає всім функціональним та нефункціональним вимогам, визначеним у технічному завданні. Використання патернів проектування State та Strategy забезпечує гнучкість та розширюваність коду.")
p.runs[0].font.size = Pt(14)
p.runs[0].font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Зберегти документ
output_file = r'C:\Users\nelox\Desktop\Курсач\Протокол_роботи.docx'
doc.save(output_file)

print("Protocol document created successfully!")
